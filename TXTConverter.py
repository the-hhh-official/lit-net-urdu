import os
from docx import Document
from docx.shared import Pt
import re

# Path to your parent folder (edit this)
parent_folder = r"F:\HHH\HHH\SNA\Books\Mustansar Hussain Tarar\raakh\results"

# Loop through each subfolder (sorted for stable order)
for folder_name in sorted(os.listdir(parent_folder)):
    folder_path = os.path.join(parent_folder, folder_name)

    # Only process if it's a directory
    if os.path.isdir(folder_path):
        doc = Document()
        print(f"Processing folder: {folder_name}")

        # Extract folder index number:
        # Matches patterns like 'footPath-1_results' -> 1
        match = re.search(r"-(\d+)(?:_|$)", folder_name)
        if match:
            folder_index = int(match.group(1))
        else:
            folder_index = 1  # Fallback if pattern not found

        # Compute starting offset for this folder
        base_page = (folder_index - 1) * 20

        # Sort txt files for consistent order
        txt_files = sorted(
            [f for f in os.listdir(folder_path) if f.lower().endswith(".txt")]
        )

        page_number = 1  # Page number within this folder

        for txt_file in txt_files:
            txt_path = os.path.join(folder_path, txt_file)

            with open(txt_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Add file content
            paragraph = doc.add_paragraph(content)
            run = paragraph.runs[0]
            run.font.size = Pt(7)

            # Global page number with offset
            full_page_num = base_page + page_number

            # Add <<<page_number>>> before page break
            footer_para = doc.add_paragraph(f"<<<{full_page_num}>>>")
            footer_run = footer_para.runs[0]
            footer_run.font.size = Pt(7)

            # Page break
            doc.add_page_break()

            page_number += 1

        # Save the document in the parent folder
        output_path = os.path.join(parent_folder, f"{folder_name}.docx")
        doc.save(output_path)
        print(f"Created: {output_path}")

print("✅ All folders processed successfully!")
